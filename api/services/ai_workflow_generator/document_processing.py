"""Document processing service for extracting text from various formats."""

from typing import List, Dict, Any
import io
import logging
from PyPDF2 import PdfReader
from docx import Document
import re

from .errors import UserInputError, ResourceError
from .logging_utils import log_operation, PerformanceLogger

logger = logging.getLogger(__name__)


class DocumentProcessingError(Exception):
    """Exception raised for document processing errors."""
    
    def __init__(self, message: str, file_type: str = None, details: str = None):
        self.message = message
        self.file_type = file_type
        self.details = details
        super().__init__(self.message)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary representation."""
        result = {"error": self.message}
        if self.file_type:
            result["file_type"] = self.file_type
        if self.details:
            result["details"] = self.details
        return result


class DocumentProcessingService:
    """Service for processing and extracting text from documents."""
    
    # Supported file types
    SUPPORTED_FORMATS = {'pdf', 'docx', 'txt', 'md'}
    
    # Maximum file size in bytes (default 10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    # Approximate characters per token (rough estimate)
    CHARS_PER_TOKEN = 4
    
    def __init__(self, max_file_size_mb: int = 10):
        """
        Initialize the document processing service.
        
        Args:
            max_file_size_mb: Maximum file size in megabytes
        """
        self.max_file_size = max_file_size_mb * 1024 * 1024
    
    async def extract_text(
        self,
        file_content: bytes,
        file_type: str
    ) -> str:
        """
        Extract text from a document.
        
        Args:
            file_content: Raw file content
            file_type: File type (pdf, docx, txt, md)
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentProcessingError: If extraction fails
        """
        # Normalize file type
        file_type = file_type.lower().strip()
        if file_type.startswith('.'):
            file_type = file_type[1:]
        
        # Validate file type
        if file_type not in self.SUPPORTED_FORMATS:
            error = UserInputError(
                message=f"Unsupported file format: {file_type}",
                details={"file_type": file_type},
                suggestions=[
                    f"Supported formats: {', '.join(self.SUPPORTED_FORMATS)}",
                    "Convert your document to a supported format"
                ]
            )
            error.log()
            raise DocumentProcessingError(
                f"Unsupported file format: {file_type}",
                file_type=file_type,
                details=f"Supported formats: {', '.join(self.SUPPORTED_FORMATS)}"
            )
        
        with log_operation(f"extract_text_{file_type}", {"file_size": len(file_content)}) as op_logger:
            try:
                # Route to appropriate extraction method
                if file_type == 'pdf':
                    text = self._extract_from_pdf(file_content)
                elif file_type == 'docx':
                    text = self._extract_from_docx(file_content)
                elif file_type in ('txt', 'md'):
                    text = self._extract_from_text(file_content)
                else:
                    raise DocumentProcessingError(
                        f"No extraction handler for file type: {file_type}",
                        file_type=file_type
                    )
                
                op_logger.log("info", f"Extracted {len(text)} characters from {file_type} document")
                return text
                
            except DocumentProcessingError:
                raise
            except Exception as e:
                logger.error(f"Failed to extract text from {file_type}: {e}", exc_info=True)
                raise DocumentProcessingError(
                    f"Failed to extract text from {file_type} file",
                    file_type=file_type,
                    details=str(e)
                )
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_content: PDF file content
            
        Returns:
            Extracted text
            
        Raises:
            DocumentProcessingError: If PDF extraction fails
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            
            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    # Log warning but continue with other pages
                    text_parts.append(f"[Error extracting page {page_num + 1}: {str(e)}]")
            
            if not text_parts:
                raise DocumentProcessingError(
                    "No text could be extracted from PDF",
                    file_type="pdf",
                    details="PDF may be empty or contain only images"
                )
            
            return "\n\n".join(text_parts)
        
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(
                "Failed to parse PDF file",
                file_type="pdf",
                details=str(e)
            )
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file content
            
        Returns:
            Extracted text
            
        Raises:
            DocumentProcessingError: If DOCX extraction fails
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                raise DocumentProcessingError(
                    "No text could be extracted from DOCX",
                    file_type="docx",
                    details="Document may be empty"
                )
            
            return "\n\n".join(text_parts)
        
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(
                "Failed to parse DOCX file",
                file_type="docx",
                details=str(e)
            )
    
    def _extract_from_text(self, file_content: bytes) -> str:
        """
        Extract text from plain text or markdown file.
        
        Args:
            file_content: Text file content
            
        Returns:
            Extracted text
            
        Raises:
            DocumentProcessingError: If text extraction fails
        """
        try:
            # Try UTF-8 first
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1
                try:
                    text = file_content.decode('latin-1')
                except UnicodeDecodeError:
                    # Last resort: ignore errors
                    text = file_content.decode('utf-8', errors='ignore')
            
            if not text.strip():
                raise DocumentProcessingError(
                    "Text file is empty",
                    file_type="txt",
                    details="No content found in file"
                )
            
            return text
        
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(
                "Failed to decode text file",
                file_type="txt",
                details=str(e)
            )
    
    def validate_file(
        self,
        file_size: int,
        file_type: str
    ) -> Dict[str, Any]:
        """
        Validate a file before processing.
        
        Args:
            file_size: Size of file in bytes
            file_type: File type
            
        Returns:
            Validation result dict with 'valid' boolean and optional 'error' message
        """
        # Normalize file type
        file_type = file_type.lower().strip()
        if file_type.startswith('.'):
            file_type = file_type[1:]
        
        # Check file type
        if file_type not in self.SUPPORTED_FORMATS:
            logger.warning(f"Unsupported file format: {file_type}")
            return {
                "valid": False,
                "error": f"Unsupported file format: {file_type}",
                "supported_formats": list(self.SUPPORTED_FORMATS)
            }
        
        # Check file size
        if file_size > self.max_file_size:
            max_mb = self.max_file_size / (1024 * 1024)
            actual_mb = file_size / (1024 * 1024)
            logger.warning(f"File size ({actual_mb:.2f}MB) exceeds maximum ({max_mb:.2f}MB)")
            
            # Log as resource error
            error = ResourceError(
                message=f"File size exceeds maximum allowed size",
                resource_type="file_size",
                details={
                    "max_size_mb": max_mb,
                    "actual_size_mb": actual_mb
                },
                suggestions=[
                    f"Maximum file size is {max_mb:.2f}MB",
                    "Reduce the file size or split into multiple files"
                ]
            )
            error.log()
            
            return {
                "valid": False,
                "error": f"File size ({actual_mb:.2f}MB) exceeds maximum allowed size ({max_mb:.2f}MB)",
                "max_size_mb": max_mb,
                "actual_size_mb": actual_mb
            }
        
        # Check for zero-size file
        if file_size == 0:
            logger.warning("Empty file uploaded")
            return {
                "valid": False,
                "error": "File is empty (0 bytes)"
            }
        
        logger.debug(f"File validation passed: {file_type}, {file_size} bytes")
        return {"valid": True}
    
    def chunk_text(
        self,
        text: str,
        max_tokens: int = 8000
    ) -> List[str]:
        """
        Chunk large text into smaller pieces for LLM processing.
        
        Uses a simple character-based approach with overlap to maintain context.
        Tries to break at paragraph boundaries when possible.
        
        Args:
            text: Text to chunk
            max_tokens: Maximum tokens per chunk (approximate)
            
        Returns:
            List of text chunks
        """
        if not text:
            return []
        
        # Convert tokens to approximate character count
        max_chars = max_tokens * self.CHARS_PER_TOKEN
        
        # If text is small enough, return as single chunk
        if len(text) <= max_chars:
            return [text]
        
        chunks = []
        
        # Split into paragraphs first
        paragraphs = re.split(r'\n\s*\n', text)
        
        current_chunk = []
        current_length = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            para_length = len(paragraph)
            
            # If single paragraph exceeds max, split it
            if para_length > max_chars:
                # Save current chunk if it has content
                if current_chunk:
                    chunks.append("\n\n".join(current_chunk))
                    current_chunk = []
                    current_length = 0
                
                # Split long paragraph by sentences
                sentences = re.split(r'(?<=[.!?])\s+', paragraph)
                temp_chunk = []
                temp_length = 0
                
                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue
                    
                    sent_length = len(sentence)
                    
                    if temp_length + sent_length > max_chars:
                        if temp_chunk:
                            chunks.append(" ".join(temp_chunk))
                        temp_chunk = [sentence]
                        temp_length = sent_length
                    else:
                        temp_chunk.append(sentence)
                        temp_length += sent_length + 1  # +1 for space
                
                if temp_chunk:
                    chunks.append(" ".join(temp_chunk))
            
            # If adding this paragraph would exceed max, save current chunk
            elif current_length + para_length > max_chars:
                if current_chunk:
                    chunks.append("\n\n".join(current_chunk))
                current_chunk = [paragraph]
                current_length = para_length
            
            # Otherwise, add to current chunk
            else:
                current_chunk.append(paragraph)
                current_length += para_length + 2  # +2 for \n\n
        
        # Add remaining chunk
        if current_chunk:
            chunks.append("\n\n".join(current_chunk))
        
        return chunks
