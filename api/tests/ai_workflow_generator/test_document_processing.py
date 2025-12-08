"""Unit tests for Document Processing Service."""

import pytest
from io import BytesIO
from PyPDF2 import PdfWriter
from docx import Document

from api.services.ai_workflow_generator.document_processing import (
    DocumentProcessingService,
    DocumentProcessingError
)


@pytest.fixture
def doc_service():
    """Create a DocumentProcessingService instance."""
    return DocumentProcessingService(max_file_size_mb=10)


@pytest.fixture
def sample_text():
    """Sample text content for testing."""
    return "This is a test document.\n\nIt has multiple paragraphs.\n\nAnd some content."


@pytest.fixture
def sample_pdf_bytes():
    """Create a simple PDF file in memory."""
    # Create a minimal valid PDF manually
    # This is a minimal PDF with text that PyPDF2 can read
    pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/Resources <<
/Font <<
/F1 <<
/Type /Font
/Subtype /Type1
/BaseFont /Helvetica
>>
>>
>>
/MediaBox [0 0 612 792]
/Contents 4 0 R
>>
endobj
4 0 obj
<<
/Length 55
>>
stream
BT
/F1 12 Tf
100 750 Td
(This is a test PDF document.) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000317 00000 n 
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
421
%%EOF
"""
    return pdf_content


@pytest.fixture
def sample_docx_bytes():
    """Create a simple DOCX file in memory."""
    doc = Document()
    doc.add_paragraph("This is a test DOCX document.")
    doc.add_paragraph("It has multiple paragraphs.")
    doc.add_paragraph("And some content.")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


class TestDocumentProcessingError:
    """Tests for DocumentProcessingError class."""
    
    def test_error_initialization(self):
        """Test DocumentProcessingError can be initialized."""
        error = DocumentProcessingError(
            "Test error",
            file_type="pdf",
            details="Test details"
        )
        
        assert error.message == "Test error"
        assert error.file_type == "pdf"
        assert error.details == "Test details"
    
    def test_error_to_dict(self):
        """Test DocumentProcessingError can be converted to dict."""
        error = DocumentProcessingError(
            "Test error",
            file_type="pdf",
            details="Test details"
        )
        
        result = error.to_dict()
        
        assert result["error"] == "Test error"
        assert result["file_type"] == "pdf"
        assert result["details"] == "Test details"


class TestDocumentProcessingService:
    """Tests for DocumentProcessingService class."""
    
    def test_initialization(self):
        """Test service can be initialized with custom max file size."""
        service = DocumentProcessingService(max_file_size_mb=5)
        assert service.max_file_size == 5 * 1024 * 1024
    
    def test_validate_file_valid(self, doc_service):
        """Test file validation with valid file."""
        result = doc_service.validate_file(1024, "pdf")
        assert result["valid"] is True
    
    def test_validate_file_unsupported_format(self, doc_service):
        """Test file validation with unsupported format."""
        result = doc_service.validate_file(1024, "exe")
        assert result["valid"] is False
        assert "unsupported" in result["error"].lower()
        assert "supported_formats" in result
    
    def test_validate_file_too_large(self, doc_service):
        """Test file validation with file too large."""
        large_size = 20 * 1024 * 1024  # 20MB
        result = doc_service.validate_file(large_size, "pdf")
        assert result["valid"] is False
        assert "exceeds" in result["error"].lower()
        assert "max_size_mb" in result
        assert "actual_size_mb" in result
    
    def test_validate_file_empty(self, doc_service):
        """Test file validation with empty file."""
        result = doc_service.validate_file(0, "pdf")
        assert result["valid"] is False
        assert "empty" in result["error"].lower()
    
    def test_validate_file_with_dot_extension(self, doc_service):
        """Test file validation with extension including dot."""
        result = doc_service.validate_file(1024, ".pdf")
        assert result["valid"] is True
    
    @pytest.mark.asyncio
    async def test_extract_text_from_txt(self, doc_service, sample_text):
        """Test text extraction from plain text file."""
        file_content = sample_text.encode('utf-8')
        result = await doc_service.extract_text(file_content, "txt")
        
        assert result == sample_text
        assert "test document" in result
    
    @pytest.mark.asyncio
    async def test_extract_text_from_md(self, doc_service):
        """Test text extraction from markdown file."""
        markdown_content = "# Heading\n\nSome **bold** text."
        file_content = markdown_content.encode('utf-8')
        result = await doc_service.extract_text(file_content, "md")
        
        assert "Heading" in result
        assert "bold" in result
    
    @pytest.mark.asyncio
    async def test_extract_text_from_txt_latin1(self, doc_service):
        """Test text extraction from latin-1 encoded file."""
        text = "Café résumé"
        file_content = text.encode('latin-1')
        result = await doc_service.extract_text(file_content, "txt")
        
        assert "Caf" in result  # Should handle encoding
    
    @pytest.mark.asyncio
    async def test_extract_text_from_pdf(self, doc_service, sample_pdf_bytes):
        """Test text extraction from PDF file."""
        result = await doc_service.extract_text(sample_pdf_bytes, "pdf")
        
        assert "test PDF document" in result
    
    @pytest.mark.asyncio
    async def test_extract_text_from_docx(self, doc_service, sample_docx_bytes):
        """Test text extraction from DOCX file."""
        result = await doc_service.extract_text(sample_docx_bytes, "docx")
        
        assert "test DOCX document" in result
        assert "multiple paragraphs" in result
    
    @pytest.mark.asyncio
    async def test_extract_text_unsupported_format(self, doc_service):
        """Test text extraction with unsupported format."""
        with pytest.raises(DocumentProcessingError) as exc_info:
            await doc_service.extract_text(b"content", "exe")
        
        assert "unsupported" in str(exc_info.value).lower()
        assert exc_info.value.file_type == "exe"
    
    @pytest.mark.asyncio
    async def test_extract_text_empty_txt(self, doc_service):
        """Test text extraction from empty text file."""
        with pytest.raises(DocumentProcessingError) as exc_info:
            await doc_service.extract_text(b"", "txt")
        
        assert "empty" in str(exc_info.value).lower()
    
    @pytest.mark.asyncio
    async def test_extract_text_corrupted_pdf(self, doc_service):
        """Test text extraction from corrupted PDF."""
        with pytest.raises(DocumentProcessingError) as exc_info:
            await doc_service.extract_text(b"not a pdf", "pdf")
        
        assert exc_info.value.file_type == "pdf"
        assert exc_info.value.details is not None
    
    @pytest.mark.asyncio
    async def test_extract_text_corrupted_docx(self, doc_service):
        """Test text extraction from corrupted DOCX."""
        with pytest.raises(DocumentProcessingError) as exc_info:
            await doc_service.extract_text(b"not a docx", "docx")
        
        assert exc_info.value.file_type == "docx"
        assert exc_info.value.details is not None
    
    def test_chunk_text_small_text(self, doc_service):
        """Test chunking with text smaller than max tokens."""
        text = "This is a small text."
        chunks = doc_service.chunk_text(text, max_tokens=1000)
        
        assert len(chunks) == 1
        assert chunks[0] == text
    
    def test_chunk_text_empty(self, doc_service):
        """Test chunking with empty text."""
        chunks = doc_service.chunk_text("", max_tokens=1000)
        assert chunks == []
    
    def test_chunk_text_large_text(self, doc_service):
        """Test chunking with text larger than max tokens."""
        # Create text that's definitely larger than max
        paragraphs = ["This is paragraph {}.".format(i) for i in range(100)]
        text = "\n\n".join(paragraphs)
        
        chunks = doc_service.chunk_text(text, max_tokens=50)  # Small max for testing
        
        assert len(chunks) > 1
        # Verify all chunks are within size limit
        max_chars = 50 * doc_service.CHARS_PER_TOKEN
        for chunk in chunks:
            assert len(chunk) <= max_chars * 1.5  # Allow some overflow for paragraph boundaries
    
    def test_chunk_text_preserves_content(self, doc_service):
        """Test that chunking preserves all content."""
        text = "Paragraph 1.\n\nParagraph 2.\n\nParagraph 3."
        chunks = doc_service.chunk_text(text, max_tokens=10)
        
        # Rejoin chunks and verify content is preserved
        rejoined = "\n\n".join(chunks)
        # Check that key content is present
        assert "Paragraph 1" in rejoined
        assert "Paragraph 2" in rejoined
        assert "Paragraph 3" in rejoined
    
    def test_chunk_text_long_paragraph(self, doc_service):
        """Test chunking with a very long paragraph."""
        # Create a long paragraph with sentences
        sentences = ["This is sentence {}.".format(i) for i in range(50)]
        text = " ".join(sentences)
        
        chunks = doc_service.chunk_text(text, max_tokens=50)
        
        assert len(chunks) > 1
        # Verify chunks contain content
        for chunk in chunks:
            assert len(chunk) > 0
    
    def test_chunk_text_with_multiple_paragraphs(self, doc_service):
        """Test chunking respects paragraph boundaries."""
        paragraphs = [
            "First paragraph with some content.",
            "Second paragraph with more content.",
            "Third paragraph with even more content."
        ]
        text = "\n\n".join(paragraphs)
        
        chunks = doc_service.chunk_text(text, max_tokens=100)
        
        # Should have chunks
        assert len(chunks) >= 1
        # First chunk should contain first paragraph
        assert "First paragraph" in chunks[0]
    
    @pytest.mark.asyncio
    async def test_extract_text_normalizes_file_type(self, doc_service, sample_text):
        """Test that file type is normalized (uppercase, with dot)."""
        file_content = sample_text.encode('utf-8')
        
        # Test with uppercase
        result1 = await doc_service.extract_text(file_content, "TXT")
        assert result1 == sample_text
        
        # Test with dot
        result2 = await doc_service.extract_text(file_content, ".txt")
        assert result2 == sample_text
        
        # Test with both
        result3 = await doc_service.extract_text(file_content, ".TXT")
        assert result3 == sample_text
    
    def test_validate_file_normalizes_file_type(self, doc_service):
        """Test that file type is normalized in validation."""
        # Test with uppercase
        result1 = doc_service.validate_file(1024, "PDF")
        assert result1["valid"] is True
        
        # Test with dot
        result2 = doc_service.validate_file(1024, ".pdf")
        assert result2["valid"] is True
        
        # Test with both
        result3 = doc_service.validate_file(1024, ".PDF")
        assert result3["valid"] is True
    
    @pytest.mark.asyncio
    async def test_extract_text_from_docx_with_table(self, doc_service):
        """Test text extraction from DOCX with tables."""
        doc = Document()
        doc.add_paragraph("Text before table.")
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"
        
        doc.add_paragraph("Text after table.")
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        result = await doc_service.extract_text(buffer.read(), "docx")
        
        assert "Text before table" in result
        assert "Cell 1" in result
        assert "Cell 4" in result
        assert "Text after table" in result
